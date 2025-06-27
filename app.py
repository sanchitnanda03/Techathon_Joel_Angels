# -*- coding: utf-8 -*-
"""
Joel's Angels - AI Legal Contracts
A professional, AI-powered legal contract generation and translation platform
"""

import os
import time
import json
from datetime import datetime, date
from io import BytesIO
import streamlit as st
from streamlit_option_menu import option_menu
from docx import Document
from dotenv import load_dotenv
from langchain.chat_models import AzureChatOpenAI
from langchain.chains import RetrievalQA
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import Docx2txtLoader
from langchain.schema import HumanMessage
from langsmith import traceable
import language_tool_python
import textstat
from utils import load_contracts_from_folder, build_vector_store, load_docs_from_folder

import streamlit as st
import os

os.environ["LANGCHAIN_API_KEY"] = st.secrets["LANGCHAIN_API_KEY"]

load_dotenv()
os.environ["LANGCHAIN_TRACING_V2"] = "true"
# os.environ["LANGCHAIN_API_KEY"] = os.getenv("LANGCHAIN_API_KEY")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
EMBED_PATH = "embeddings"
INDEX_FILE = os.path.join(EMBED_PATH, "index.faiss")
if not os.path.exists(INDEX_FILE):
    st.info("🔄 First-time setup: creating vector store from templates...")
    docs = load_contracts_from_folder("contract_templates")
    vectordb = build_vector_store(docs, EMBED_PATH)
else:
    vectordb = FAISS.load_local(
        EMBED_PATH,
        HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2"),
        allow_dangerous_deserialization=True
    )
    retriever = vectordb.as_retriever(search_kwargs={"k": 3})

@st.cache_resource
def build_retriever():
    from pathlib import Path
    folder = Path("reference_docs")
    docs = []
    for f in folder.glob("*.docx"):
        docs.extend(Docx2txtLoader(str(f)).load())
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=80)
    splits = text_splitter.split_documents(docs)
    embedding = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2", model_kwargs={"device": "cpu"})
    vectordb = FAISS.from_documents(splits, embedding)
    return vectordb.as_retriever()

retriever1 = build_retriever()

llm = AzureChatOpenAI(
    openai_api_key=OPENAI_API_KEY,
    azure_endpoint="https://openaiqc.gep.com/techathon/openai/deployments/gpt-4o-mini/chat/completions?api-version=2025-01-01-preview",
    openai_api_version="2025-01-01-preview",
    deployment_name="gpt-4o-mini",
    temperature=0.3
)

qa_chain = RetrievalQA.from_chain_type(
    llm=llm,
    retriever=retriever,
    return_source_documents=False
)

qa_chain1 = RetrievalQA.from_chain_type(
    llm=llm,
    retriever=retriever1,
    return_source_documents=False
)

@traceable(name="generate_contract_template_interaction")
def generate_contract_template(query):
    return qa_chain.run(query)

@traceable(name="translate_doc_interaction")
def translate_doc(doc_text: str, target_lang: str) -> str:
    prompt = (
        f"Translate the following contract into {target_lang.upper()} using similar legal structure and terminology "
        f"as seen in the reference documents:\n\n{doc_text}"
    )
    return qa_chain1.run(prompt)

# Page configuration
st.set_page_config(
    page_title="Joel's Angels - AI Legal Contracts",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for styling
def load_css():
    st.markdown("""
    <style>
        /* Hide default Streamlit elements */
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        
        /* Feature card styling - LIGHT THEME */
        .feature-card {
            background: white;
            padding: 2.5rem;
            border-radius: 15px;
            box-shadow: 0 8px 25px rgba(0, 0, 0, 0.08);
            margin: 1rem 0;
            border-left: 5px solid #0181B0;
            transition: all 0.3s ease;
        }
        
        .feature-card h3 {
            color: #1F2937 !important;
            font-weight: bold;
            margin-bottom: 1rem;
            font-size: 1.3rem;
        }
        
        .feature-card p {
            color: #6B7280 !important;
            line-height: 1.6;
            margin: 0;
            font-size: 1rem;
        }
        
        .feature-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 12px 35px rgba(0, 0, 0, 0.12);
        }
        
        /* Main content styling - LIGHT THEME */
        .main-header {
            background: linear-gradient(90deg, #0181B0 0%, #3793BA 100%);
            padding: 2rem;
            border-radius: 15px;
            margin-bottom: 2rem;
            color: white;
            text-align: center;
            box-shadow: 0 4px 15px rgba(99, 102, 241, 0.2);
        }
        
        .hero-section {
            background: linear-gradient(135deg, #F3F4F6 0%, #E5E7EB 100%);
            padding: 4rem 2rem;
            border-radius: 20px;
            margin: 2rem 0;
            color: #1F2937;
            text-align: center;
            position: relative;
            overflow: hidden;
            border: 1px solid #E5E7EB;
        }
        
        .hero-section::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: url('data:image/svg+xml,<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100"><text y="50" font-size="8" fill="rgba(99,102,241,0.1)">📄</text></svg>') repeat;
            animation: float 6s ease-in-out infinite;
        }
        
        @keyframes float {
            0%, 100% { transform: translateY(0px); }
            50% { transform: translateY(-10px); }
        }
        
        .hero-content {
            position: relative;
            z-index: 2;
        }
        
        .contract-form {
            background: #F9FAFB;
            padding: 2.5rem;
            border-radius: 15px;
            border: 2px solid #E5E7EB;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.05);
        }
        
        .preview-panel {
            background: white;
            padding: 2.5rem;
            border-radius: 15px;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.08);
            border: 1px solid #E5E7EB;
            min-height: 400px;
        }
        
        .ai-assistant {
            position: fixed;
            bottom: 30px;
            right: 30px;
            background: linear-gradient(45deg, #0181B0, #3793BA);
            color: white;
            padding: 1.5rem;
            border-radius: 50px;
            box-shadow: 0 8px 25px rgba(99, 102, 241, 0.3);
            z-index: 1000;
            cursor: pointer;
            animation: pulse 2s infinite;
            display: flex;
            align-items: center;
            gap: 0.5rem;
            font-weight: bold;
        }
        
        @keyframes pulse {
            0% { transform: scale(1); box-shadow: 0 8px 25px rgba(99, 102, 241, 0.3); }
            50% { transform: scale(1.05); box-shadow: 0 12px 35px rgba(99, 102, 241, 0.4); }
            100% { transform: scale(1); box-shadow: 0 8px 25px rgba(99, 102, 241, 0.3); }
        }
        
        .security-badge {
            background: linear-gradient(45deg, #10B981, #059669);
            color: white;
            padding: 0.75rem 1.5rem;
            border-radius: 25px;
            font-weight: bold;
            display: inline-block;
            margin: 0.5rem 0;
            box-shadow: 0 4px 15px rgba(16, 185, 129, 0.3);
            animation: glow 2s ease-in-out infinite alternate;
        }
        
        @keyframes glow {
            from { box-shadow: 0 4px 15px rgba(16, 185, 129, 0.3); }
            to { box-shadow: 0 4px 25px rgba(16, 185, 129, 0.5); }
        }
        
        .risk-indicator {
            display: inline-block;
            padding: 0.5rem 1rem;
            border-radius: 20px;
            font-size: 0.9rem;
            font-weight: bold;
            box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1);
        }
        
        .risk-low { background: #D1FAE5; color: #065F46; }
        .risk-medium { background: #FEF3C7; color: #92400E; }
        .risk-high { background: #FEE2E2; color: #991B1B; }
        
        .stButton > button {
            background: linear-gradient(45deg, #0181B0, #3793BA);
            color: white;
            border: none;
            border-radius: 25px;
            padding: 1rem 2rem;
            font-weight: bold;
            transition: all 0.3s ease;
            box-shadow: 0 4px 15px rgba(99, 102, 241, 0.3);
        }
        
        .stButton > button:hover {
            transform: translateY(-3px);
            box-shadow: 0 8px 25px rgba(99, 102, 241, 0.4);
        }
        
        .stButton > button:active {
            transform: translateY(-1px);
        }
        
        /* Form field styling - LIGHT THEME */
        .stTextInput > div > div > input,
        .stSelectbox > div > div > select,
        .stDateInput > div > div > input {
            border-radius: 10px;
            border: 2px solid #E5E7EB;
            transition: all 0.3s ease;
            background: white;
        }
        
        .stTextInput > div > div > input:focus,
        .stSelectbox > div > div > select:focus,
        .stDateInput > div > div > input:focus {
            border-color: #0181B0;
            box-shadow: 0 0 0 3px rgba(99, 102, 241, 0.1);
        }
        
        /* Progress bar styling - LIGHT THEME */
        .stProgress > div > div > div > div {
            background: linear-gradient(90deg, #0181B0, #3793BA);
        }
        
        /* Chat message styling */
        .stChatMessage {
            border-radius: 15px;
            margin: 1rem 0;
        }
        
        /* File upload styling - LIGHT THEME */
        .stFileUploader > div {
            border: 2px dashed #0181B0;
            border-radius: 15px;
            padding: 2rem;
            text-align: center;
            transition: all 0.3s ease;
            background: #25252C;
        }
        
        .stFileUploader > div:hover {
            border-color: #3793BA;
            background: rgba(99, 102, 241, 0.05);
        }
        
        /* Header styling - LIGHT THEME */
        .stSelectbox > div > div > select {
            background: rgba(255, 255, 255, 0.9) !important;
            color: #1F2937 !important;
            border: 1px solid rgba(255, 255, 255, 0.3) !important;
            border-radius: 20px !important;
            padding: 0.3rem 0.8rem !important;
            font-size: 0.8rem !important;
            min-width: 80px !important;
        }
        
        .stSelectbox > div > div > select:focus {
            border-color: #0181B0 !important;
            box-shadow: 0 0 0 2px rgba(99, 102, 241, 0.3) !important;
        }
        
        .stSelectbox > div > div > select option {
            background: white !important;
            color: #1F2937 !important;
        }
        
        .stSelectbox > label {
            display: none !important;
        }
        
        .stButton > button[data-testid="baseButton-secondary"] {
            background: rgba(255, 255, 255, 0.9) !important;
            color: #1F2937 !important;
            border: 1px solid rgba(255, 255, 255, 0.3) !important;
            border-radius: 50% !important;
            width: 35px !important;
            height: 35px !important;
            padding: 0 !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            font-size: 1rem !important;
            transition: all 0.3s ease !important;
            min-width: 35px !important;
        }
        
        .stButton > button[data-testid="baseButton-secondary"]:hover {
            background: rgba(255, 255, 255, 1) !important;
            transform: scale(1.1) !important;
        }
        
        /* Navigation button styling - LIGHT THEME */
        .nav-button {
            background: linear-gradient(45deg, #F3F4F6, #E5E7EB);
            color: #1F2937;
            border: none;
            border-radius: 10px;
            padding: 0.75rem 1.5rem;
            font-weight: 500;
            transition: all 0.3s ease;
            cursor: pointer;
            margin: 0.25rem;
        }
        
        .nav-button:hover {
            background: linear-gradient(45deg, #0181B0, #3793BA);
            color: white;
            transform: translateY(-2px);
            box-shadow: 0 4px 15px rgba(99, 102, 241, 0.3);
        }
        
        .nav-button.active {
            background: linear-gradient(45deg, #0181B0, #3793BA);
            color: white;
            font-weight: bold;
        }
        
        /* Page background */
        .main .block-container {
            background: #FAFAFA;
            padding-top: 2rem;
            padding-bottom: 2rem;
        }
        .main {
            background-color: #121212;
            color: #F0F0F0;
        }

        
        .stTextInput > div > div > input {
            background-color: #2e2e2e;  /* Dark grey background */
            color: white;               /* White text for contrast */
            border: 1px solid #555;     /* Optional: subtle border */
        }

        </style>

    """, unsafe_allow_html=True)

# Initialize session state
if 'contract_data' not in st.session_state:
    st.session_state.contract_data = {}
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'current_page' not in st.session_state:
    st.session_state.current_page = "Home"

# Load CSS
load_css()

# Custom Header Component
def create_header():
    # Create header with columns
    header_col1, header_col2 = st.columns([3, 7])
    
    with header_col1:
        # Logo section
        st.markdown("""
        <div style="display: flex; align-items: center; gap: 0.5rem;">
            <div style="font-size: 2rem; background: linear-gradient(45deg, #FCD34D, #F59E0B); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;">⚖️</div>
            <div>
                <h1 style="color: white; font-size: 1.5rem; font-weight: bold; margin: 0;">Joel's Angels</h1>
                <p style="color: #E0E7FF; font-size: 0.8rem; margin: 0; opacity: 0.9;">AI-Powered Legal Contracts</p>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with header_col2:
        # Empty space for centering
        st.markdown("")

# Home page
def home_page():
    st.markdown("""
    <div class="hero-section">
        <div class="hero-content">
            <h1 style="font-size: 3rem; margin-bottom: 1rem;">AI-Powered Legal Contracts</h1>
            <h2 style="font-size: 1.5rem; margin-bottom: 2rem;">Precision-Tailored in 60 Seconds</h2>
            <p style="font-size: 1.3rem; margin: 1rem 0; opacity: 0.9;">Court-Ready Documents • Enterprise-Grade Security • Zero Legal Jargon</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # CTA Buttons
    col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
    
    with col2:
        if st.button("🚀 Generate Template", use_container_width=True, key="hero_generate"):
            st.session_state.current_page = "Generate Contract"
            st.rerun()
    
    with col3:
        if st.button("🌐 Translate Document", use_container_width=True, key="hero_translate"):
            st.session_state.current_page = "Translate Document"
            st.rerun()
    
    # Feature cards
    st.markdown("## ✨ Key Features")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="feature-card">
            <h3>⚡ Lightning Fast</h3>
            <p>Generate professional contracts in under 60 seconds with our advanced AI technology.</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <h3>🔒 Secure & Compliant</h3>
            <p>Enterprise-grade encryption ensures your documents remain confidential and legally sound.</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="feature-card">
            <h3>🌍 Multi-Language</h3>
            <p>Translate legal documents across multiple languages while preserving legal accuracy.</p>
        </div>
        """, unsafe_allow_html=True)
    
    # Navigation section
    st.markdown("## 🧭 Quick Navigation")
    nav_col1, nav_col2, nav_col3, nav_col4 = st.columns(4)
    
    with nav_col1:
        if st.button("🤖 AI Assistant", use_container_width=True, key="nav_ai_home"):
            st.session_state.current_page = "AI Assistant"
            st.rerun()
    
    with nav_col2:
        if st.button("ℹ️ About Us", use_container_width=True, key="nav_about_home"):
            st.session_state.current_page = "About"
            st.rerun()
    
    with nav_col3:
        if st.button("📞 Contact", use_container_width=True, key="nav_contact_home"):
            st.session_state.current_page = "Contact"
            st.rerun()
    
    with nav_col4:
        if st.button("📄 Generate Template", use_container_width=True, key="nav_generate_home"):
            st.session_state.current_page = "Generate Contract"
            st.rerun()

# Contract Generation Page
def contract_generation_page():
    st.markdown("""
    <div class="main-header" style="background: linear-gradient(90deg, #0181B0 0%, #3793BA 100%); ">
        <h1>📄 Generate Legal Contract Language Template</h1>
        <p>Create professional, formal contract language templates in minutes</p>
    </div>
    """, unsafe_allow_html=True)
    
    with open("contract_clause_keyword.json", "r") as f:
        required_keywords = json.load(f)

 
    def calculate_actqm(contract_text, contract_type, penalty_factor=0.5):
        # Clause Coverage Score (CCS)
        must_have_keywords = required_keywords.get(contract_type, [])
        found_keywords = [kw for kw in must_have_keywords if kw.lower() in contract_text.lower()]
        ccs = len(found_keywords) / len(must_have_keywords) if must_have_keywords else 0
    
        # Formality Score (FS) - less strict penalty
        tool = language_tool_python.LanguageTool('en-US')
        matches = tool.check(contract_text)
        grammar_issues = len(matches)
        sentence_count = max(textstat.sentence_count(contract_text), 1)
        fs = 1 - penalty_factor * (grammar_issues / sentence_count)
        fs = max(fs, 0)  # ensure FS not below zero
    
        # ACTQM Score
        actqm = ((2 * ccs + fs) / 3) * 100
        return {
            "CCS": round(ccs, 2),
            "FS": round(fs, 2),
            "ACTQM": round(actqm, 2),
            "Keywords Found": found_keywords,
            "Total Required": len(must_have_keywords),
            "Grammar Issues": grammar_issues,
            "Sentences": sentence_count
        }

    if "history" not in st.session_state:
        st.session_state.history = []

    # Back to Home button
    if st.button("🏠 Back to Home", key="back_home_generate"):
        st.session_state.current_page = "Home"
        st.rerun()

    result = None
    st.markdown("""
    <div class="contract-form" style="background: #222222; border: None;">
        <h3>Contract Template Details</h3>
    </div>
    """, unsafe_allow_html=True)
    
    # Contract category with icons
    categories = {
        "🏠 NDA": "Non-Disclosure Agreement",
        "🤝 MSA": "Master Services Agreement", 
        "💼 Employment": "Employment Agreement",
        "💰 Sponsorship": "Sponsorship Agreement",
    }

    client_name = st.selectbox("Select Client Folder", sorted(os.listdir("client_metadata")))
    category = st.selectbox("Contract Category", list(categories.keys()))
    contract_type = categories[category]
    
    docs = load_docs_from_folder(os.path.join("client_metadata", client_name))
    if docs and len(docs) > 0:
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        vectordb = FAISS.from_documents(docs, embeddings)
        retriever2 = vectordb.as_retriever(search_kwargs={"k": 3})
        qa_chain2 = RetrievalQA.from_chain_type(
            llm=llm, 
            retriever=retriever2
        )
    else:
        st.error(f"No documents found in 'client_metadata/{client_name}'. Please check the folder or upload valid files.")
        qa_chain2 = None  # Optional: prevent future crashes if qa_chain2 is used


    # Effective date
    effective_date = st.date_input("Effective Date", value=date.today())
    
    # Advanced options
    with st.expander("⚙️ Advanced Options"):
        jurisdiction = st.selectbox("Jurisdiction", ["United States", "United Kingdom", "Canada", "Australia", "Other"])
        governing_law = st.selectbox("Governing Law", ["Common Law", "Civil Law", "Mixed Jurisdiction"])
        contract_value = st.number_input("Contract Value ($)", min_value=0, value=10000)
    
    # Generate button
    if st.button("🚀 Generate Contract Language Template", use_container_width=True, type="primary") and category and effective_date:
        with st.spinner("Generating your contract..."):
            # Simulate contract generation with building blocks animation
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            stages = ["Drafting...", "Reviewing...", "Finalizing..."]
            for i, stage in enumerate(stages):
                status_text.text(stage)
                for j in range(33):
                    time.sleep(0.02)
                    progress_bar.progress((i * 33 + j + 1) * 100 // 100)
            
            prompt = "Analyze the client's contract style. Extract common clause types, writing tone, preferred keywords, and jurisdiction references. Return them in bullet points."
            response = qa_chain2.run(prompt)
            query = f"Create a detailed contract template for {contract_type}.Make sure that the generated template is based on {response}. Make sure it is formal, general-purpose, and does not include any party names.Value of the contract is {contract_value}. Jurisdiction is {jurisdiction}. Governing law is {governing_law}. Effective date is {effective_date}."
            st.session_state.history.append(("user_gen", query))
            result = generate_contract_template(query)
            st.session_state.history.append(("ai_gen", result))
            st.success("Contract generated successfully!")

    if result:
        st.text_area("📄 Contract Template Preview", result, height=400)
        
        filename = f"{contract_type.title().replace(' ', '_')}_template.txt"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(result)

        with open(filename, "rb") as f:
            st.download_button("📥 Download Template as Text File", f, file_name=filename, mime="text/plain", use_container_width=True)    
        # Human review option
        if st.button("👨‍💼 Request Human Review", use_container_width=True):
            st.info("Human review request sent! A legal expert will review your contract within 24 hours.")
        
        st.subheader("📊 ACTQM: Automated Quality Evaluation")
        metrics = calculate_actqm(result, contract_type.strip())
        st.markdown(f"""
        <style>
            .metric-card {{
                background: linear-gradient(90deg, #0181B0 0%, #3793BA 100%);
                padding: 1rem 1.5rem;
                border-radius: 12px;
                box-shadow: 0 8px 20px rgba(0, 0, 0, 0.15);
                color: white;
                font-family: 'Segoe UI', sans-serif;
                margin-bottom: 1.5rem;
            }}
            .metric-card li {{
                margin: 0.5rem 0;
                font-size: 1.1rem;
                list-style: none;
            }}
            .metric-card li::before {{
                content: "📊 ";
                margin-right: 0.5rem;
            }}
            .metric-card li:last-child::before {{
                content: "✅ ";
            }}
            .metric-card code {{
                background: rgba(255, 255, 255, 0.1);
                padding: 2px 6px;
                border-radius: 4px;
            }}
        </style>

        <div class="metric-card">
            <ul>
                <li><strong>Clause Coverage Score (CCS)</strong>: <code>{metrics['CCS']}</code></li>
                <li><strong>Formality Score (FS)</strong>: <code>{metrics['FS']}</code></li>
                <li><strong>ACTQM</strong>: <code>{metrics['ACTQM']}</code></li>
            </ul>
        </div>
    """, unsafe_allow_html=True)


    # Feedback and history
    for i, (role, msg) in enumerate(st.session_state.history):
        if role == "ai_gen":
            feedback_key = f"feedback_{i}"
            feedback = st.radio(
                "Was this response helpful?",
                ("👍 Yes", "👎 No"),
                key=feedback_key
            )
            if feedback == "👍 Yes":
                st.success("✅ Thanks for your feedback!")
            elif feedback == "👎 No":
                st.warning("❗ We appreciate your feedback and will improve.")
                client_feedback = st.text_input("Please provide the changes you require:", key=f"feedback_text_{i}")
                if st.button("Submit changes", key=f"submit_changes_{i}") and client_feedback:
                    modified_query = f"Make the changes to {msg} based on the feedback: {client_feedback}"
                    modified_result = generate_contract_template(modified_query)
                    st.session_state.history.append(("ai_gen", modified_result))
                    st.subheader("📑 Modified Contract Template")
                    st.text_area("Modified Template", modified_result, height=400)
                    filename = f"{contract_type.title().replace(' ', '_')}_template.txt"
                    with open(filename, "w", encoding="utf-8") as f:
                        f.write(modified_result)
                    with open(filename, "rb") as f:
                        st.download_button("📥 Download Template as Text File", f, file_name=filename, mime="text/plain", use_container_width=True)
                    if st.button("👨‍💼 Request Human Review", use_container_width=True):
                        st.info("Human review request sent! A legal expert will review your contract within 24 hours.")
                    st.subheader("📊 ACTQM: Automated Quality Evaluation")
                    metrics = calculate_actqm(modified_result, contract_type.strip())
                    st.markdown(f"""
                        - **Clause Coverage Score (CCS)**: `{metrics['CCS']} ({len(metrics['Keywords Found'])}/{metrics['Total Required']} keywords found)`
                        - **Formality Score (FS)**: `{metrics['FS']} ({metrics['Grammar Issues']} issues in {metrics['Sentences']} sentences)`
                        - **✅ ACTQM**: `{metrics['ACTQM']}`
                        """)

# Translation Portal
def translation_page():
    st.markdown("""
    <div class="main-header">
        <h1>🌐 Document Translation</h1>
        <p>Translate legal documents with AI-powered accuracy</p>
    </div>
    """, unsafe_allow_html=True)

    if "history" not in st.session_state:
        st.session_state.history = []

    # Back to Home button
    if st.button("🏠 Back to Home", key="back_home_translate"):
        st.session_state.current_page = "Home"
        st.rerun()
    
    # File upload zone with drag-and-drop styling
    st.markdown("### 📄 Upload Document")
    uploaded_file = st.file_uploader("📄 Upload a DOCX file", type=["docx"], help="Supported format: DOCX")
    
    if uploaded_file:
        st.success(f"✅ File uploaded: {uploaded_file.name}")
        
        # Language selection with flag icons
        st.markdown("### 🌍 Select Languages")
        col1, col2 = st.columns(2)
        
        with col1:
            source_language = st.selectbox(
                "From Language",
                ["English", "Spanish", "French", "German", "Dutch"]
            )
        
        with col2:
            target_language = st.selectbox(
                "To Language", 
                ["English", "Spanish", "French", "German", "Dutch"]
            )
        
        # Translate button
        if st.button("🌐 Translate Document", use_container_width=True, type="primary"):
            with st.spinner("Translating your document..."):
                progress_bar = st.progress(0)
                for i in range(100):
                    time.sleep(0.03)
                    progress_bar.progress(i + 1)
                
                doc = Document(uploaded_file)
                full_text = "\n".join([p.text for p in doc.paragraphs])

                with st.spinner("🔄 Translating..."):
                    st.session_state.history.append(("user_trans", full_text))
                    translated = translate_doc(full_text, target_language)
                    st.session_state.history.append(("ai_trans", translated))

                st.success("✅ Translation complete!")

                # 📝 Show translated content in a text area
                st.subheader("📄 Translated Template")
                st.text_area("Review the translated document below:", translated, height=400)

                # 📄 Create translated DOCX
                translated_doc = Document()
                for para in translated.split("\n"):
                    if para.strip():
                        translated_doc.add_paragraph(para.strip())

                output = BytesIO()
                translated_doc.save(output)
                output.seek(0)

                st.download_button(
                    label="📥 Download Translated DOCX",
                    data=output,
                    file_name=f"translated_{target_language}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
        # User feedback section
        for i, (role, msg) in enumerate(st.session_state.history):
            if role == "ai_trans":
                feedback_key = f"feedback_{i}"
                feedback = st.radio(
                    "Was this response helpful?",
                    ("👍 Yes", "👎 No"),
                    key=feedback_key
                )
                if feedback == "👍 Yes":
                    st.success("✅ Thanks for your feedback!")
                elif feedback == "👎 No":
                    st.warning("❗ Please provide what changes you would want to see.")
                    client_feedback = st.text_input("Please provide the changes you require:", key=f"feedback_text_{i}")
                    if st.button("Submit changes", key=f"submit_changes_{i}") and client_feedback:
                        modified_query = f"Make the changes to: {msg} based on the feedback: {client_feedback}"
                        modified_result = qa_chain1.run(modified_query)
                        st.session_state.history.append(("ai_trans", modified_result))
                        st.subheader("📑 Modified Language Template Translation")
                        st.text_area("Modified Translation", modified_result, height=400)
                        translated_doc = Document()
                        for para in modified_result.split("\n"):
                            if para.strip():
                                translated_doc.add_paragraph(para.strip())

                        output = BytesIO()
                        translated_doc.save(output)
                        output.seek(0)

                        st.download_button(
                            label="📥 Download Translated DOCX",
                            data=output,
                            file_name=f"translated_{target_language}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

# AI Assistant
def ai_assistant_page():
    st.markdown("""
    <div class="main-header">
        <h1>🤖 LegalMind Assistant</h1>
        <p>Your AI-powered legal guidance companion</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Back to Home button
    if st.button("🏠 Back to Home", key="back_home_ai"):
        st.session_state.current_page = "Home"
        st.rerun()
    
    # Chat interface
    st.markdown("### 💬 Chat with LegalMind")
    
    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = [
            {"role": "assistant", "content": "Hello! I'm LegalMind, your AI legal assistant. I can help you with contract generation, legal questions, and document analysis. How can I assist you today?"}
        ]
    
    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # Chat input
    if prompt := st.chat_input("Ask me anything about legal contracts..."):
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        # Display user message
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Generate AI response
        with st.chat_message("assistant"):
            with st.spinner("LegalMind is thinking..."):
                # Simulate AI response
                time.sleep(1)
                
                response = llm([HumanMessage(content=prompt)]).content              
                st.markdown(response)
                st.session_state.messages.append({"role": "assistant", "content": response})
    
    # Quick actions
    st.markdown("### ⚡ Quick Actions")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 Generate Contract", use_container_width=True):
            st.session_state.current_page = "Generate Contract"
            st.rerun()
    
    with col2:
        if st.button("🌐 Translate Document", use_container_width=True):
            st.session_state.current_page = "Translate Document"
            st.rerun()


# About page
def about_page():
    st.markdown("""
    <div class="main-header">
        <h1>ℹ️ About Joel's Angels</h1>
        <p>Revolutionizing legal document creation with AI</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Back to Home button
    if st.button("🏠 Back to Home", key="back_home_about"):
        st.session_state.current_page = "Home"
        st.rerun()
    
    st.markdown("""
    ## Our Mission
    
    Joel's Angels is dedicated to democratizing access to professional legal documents through cutting-edge AI technology. 
    We believe that everyone deserves access to high-quality, legally sound contracts without the traditional barriers of cost and complexity.
    
    ## Key Features
    
    - **AI-Powered Generation**: Create professional contracts in under 60 seconds
    - **Multi-Language Support**: Translate legal documents across multiple languages
    - **Enterprise Security**: Bank-level encryption and compliance standards
    - **Legal Accuracy**: Court-ready documents that stand up to legal scrutiny
    - **24/7 AI Assistant**: Get instant legal guidance anytime, anywhere
    
    ## Technology Stack
    
    - **Frontend**: Streamlit for rapid development and beautiful UI
    - **AI/ML**: Advanced language models for contract generation
    - **Security**: Enterprise-grade encryption and secure document handling
    - **Compliance**: Built-in legal compliance and jurisdiction handling
    
    ## Contact Information
    
    - **Email**: contact@joelsangels.com
    - **Phone**: +1 (555) 123-4567
    - **Address**: 123 Legal Tech Avenue, Innovation District
    
    ## Security & Compliance
    
    <div class="security-badge">
        🔒 SOC 2 Type II Certified
    </div>
    
    <div class="security-badge">
        🛡️ GDPR Compliant
    </div>
    
    <div class="security-badge">
        ⚖️ Legal Industry Certified
    </div>
    """, unsafe_allow_html=True)

# Contact page
def contact_page():
    st.markdown("""
    <div class="main-header">
        <h1>📞 Contact Us</h1>
        <p>Get in touch with our legal tech experts</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Back to Home button
    if st.button("🏠 Back to Home", key="back_home_contact"):
        st.session_state.current_page = "Home"
        st.rerun()
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("### 📧 Send us a message")
        
        name = st.text_input("Full Name")
        email = st.text_input("Email Address")
        subject = st.selectbox("Subject", ["General Inquiry", "Technical Support", "Business Partnership", "Legal Consultation"])
        message = st.text_area("Message", height=200)
        
        if st.button("📤 Send Message", use_container_width=True, type="primary"):
            st.success("Message sent successfully! We'll get back to you within 24 hours.")
    
    with col2:
        st.markdown("### 📍 Contact Information")
        
        st.markdown("""
        **🏢 Office Address**  
        123 Legal Tech Avenue  
        Innovation District  
        Tech City, TC 12345
        
        **📞 Phone**  
        +1 (555) 123-4567
        
        **📧 Email**  
        contact@joelsangels.com
        
        **🌐 Website**  
        www.joelsangels.com
        
        **⏰ Business Hours**  
        Monday - Friday: 9:00 AM - 6:00 PM EST  
        Saturday: 10:00 AM - 2:00 PM EST
        """)

# Main app logic
def main():
    # Create custom header
    create_header()
    
    # Page routing
    current_page = st.session_state.current_page
    
    if current_page == "Home":
        home_page()
    elif current_page == "Generate Contract":
        contract_generation_page()
    elif current_page == "Translate Document":
        translation_page()
    elif current_page == "AI Assistant":
        ai_assistant_page()
    elif current_page == "About":
        about_page()
    elif current_page == "Contact":
        contact_page()
    
    # Floating AI Assistant (always visible)
    st.markdown("""
    <div class="ai-assistant" onclick="alert('AI Assistant clicked!')">
        🦉 LegalMind
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
