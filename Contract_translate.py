import os
import streamlit as st
from docx import Document
from io import BytesIO
from langchain.chat_models import AzureChatOpenAI
from langchain.chains import RetrievalQA
from langchain_community.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langsmith import traceable
from dotenv import load_dotenv

# Load .env if using it
load_dotenv()

# 🔐 API Keys (should be stored in secrets or .env)
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"] = st.secrets.get("LANGCHAIN_API_KEY", "")
os.environ["AZURE_OPENAI_API_KEY"] = st.secrets.get("AZURE_OPENAI_API_KEY", "")

# 🔧 Azure OpenAI setup
AZURE_DEPLOYMENT_NAME = "gpt-4o-mini"
AZURE_API_BASE = "https://openaiqc.gep.com/techathon/openai/deployments/gpt-4o-mini/chat/completions?api-version=2025-01-01-preview"  # ✅ DO NOT include "/openai/..." here
AZURE_API_VERSION = "2025-01-01-preview"

# 🎛️ Streamlit UI
st.set_page_config("📝 Document Translator")
st.title("🌍 Document Translator with Reference Context + LangSmith")

# 📤 Upload section
uploaded_file = st.file_uploader("📄 Upload a DOCX file", type=["docx"])

# 🌐 Language selection
languages = {
    "German": "de",
    "Spanish": "es",
    "French": "fr",
    "Dutch": "nl"
}
selected_lang = st.selectbox("🌐 Translate to:", list(languages.keys()))

# 📚 Reference Docs for RAG
@st.cache_resource
def build_retriever():
    from pathlib import Path
    folder = Path("reference_docs")
    docs = []
    for f in folder.glob("*.docx"):
        docs.extend(Docx2txtLoader(str(f)).load())
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=80)
    splits = text_splitter.split_documents(docs)
    embedding = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    vectordb = FAISS.from_documents(splits, embedding)
    return vectordb.as_retriever()

retriever = build_retriever()

# 🤖 LLM setup
llm = AzureChatOpenAI(
    deployment_name=AZURE_DEPLOYMENT_NAME,
    openai_api_key=os.environ["AZURE_OPENAI_API_KEY"],
    azure_endpoint=AZURE_API_BASE,
    openai_api_version=AZURE_API_VERSION,
    temperature=0.3
)
qa = RetrievalQA.from_chain_type(llm=llm, retriever=retriever, return_source_documents=False)

if "history" not in st.session_state:
    st.session_state.history = []

# 🧠 Translation function
@traceable(name="translate_doc_interaction")
def translate_doc(doc_text: str, target_lang: str) -> str:
    prompt = (
        f"Translate the following contract into {target_lang.upper()} using similar legal structure and terminology "
        f"as seen in the reference documents:\n\n{doc_text}"
    )
    return qa.run(prompt)

# 🚀 Perform translation
if st.button("🔄 Translate Document") and uploaded_file and selected_lang:
    doc = Document(uploaded_file)
    full_text = "\n".join([p.text for p in doc.paragraphs])

    with st.spinner("🔄 Translating..."):
        st.session_state.history.append(("user", full_text))
        translated = translate_doc(full_text, selected_lang)
        st.session_state.history.append(("ai", translated))

    st.success("✅ Translation complete!")

    # 📝 Show translated content in a text area
    st.subheader("📄 Translated Template")
    st.text_area("Review the translated document below:", translated, height=400)

    # 📄 Create translated DOCX
    translated_doc = Document()
    for para in translated.split("\n"):
        if para.strip():
            translated_doc.add_paragraph(para.strip())

    output = BytesIO()
    translated_doc.save(output)
    output.seek(0)

    st.download_button(
        label="📥 Download Translated DOCX",
        data=output,
        file_name=f"translated_{languages[selected_lang]}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# User feedback section
for i, (role, msg) in enumerate(st.session_state.history):
    if role == "ai":
        feedback_key = f"feedback_{i}"
        feedback = st.radio(
            "Was this response helpful?",
            ("👍 Yes", "👎 No"),
            key=feedback_key
        )
        if feedback == "👍 Yes":
            st.success("✅ Thanks for your feedback!")
        elif feedback == "👎 No":
            st.warning("❗ Please provide what changes you would want to see.")
            client_feedback = st.text_input("Please provide the changes you require:", key=f"feedback_text_{i}")
            if st.button("Submit changes", key=f"submit_changes_{i}") and client_feedback:
                modified_query = f"Make the changes to: {msg} based on the feedback: {client_feedback}"
                modified_result = qa.run(modified_query)
                st.session_state.history.append(("ai", modified_result))
                st.subheader("📑 Modified Language Template Translation")
                st.text_area("Modified Translation", modified_result, height=400)
                translated_doc = Document()
                for para in modified_result.split("\n"):
                    if para.strip():
                        translated_doc.add_paragraph(para.strip())

                output = BytesIO()
                translated_doc.save(output)
                output.seek(0)

                st.download_button(
                    label="📥 Download Translated DOCX",
                    data=output,
                    file_name=f"translated_{languages[selected_lang]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
