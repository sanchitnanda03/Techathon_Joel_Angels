import os
import streamlit as st
from dotenv import load_dotenv
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.chat_models import AzureChatOpenAI
from utils.loader import load_contracts_from_folder, build_vector_store
from docx import Document

load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

st.set_page_config(page_title="📄 Contract Template Generator", layout="centered")
st.title("🤖 Contract Template Generator using Azure GPT-4o-mini + FAISS")

# Load or create Vector Store
EMBED_PATH = "embeddings"
INDEX_FILE = os.path.join(EMBED_PATH, "index.faiss")

if not os.path.exists(INDEX_FILE):
    st.info("🔄 First-time setup: creating vector store from templates...")
    docs = load_contracts_from_folder("contract_templates")
    vectordb = build_vector_store(docs, EMBED_PATH)
else:
    vectordb = FAISS.load_local(
        EMBED_PATH,
        HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2"),
        allow_dangerous_deserialization=True
    )

retriever = vectordb.as_retriever(search_kwargs={"k": 3})

llm = AzureChatOpenAI(
    openai_api_key=OPENAI_API_KEY,
    azure_endpoint="https://openaiqc.gep.com/techathon/openai/deployments/gpt-4o-mini/chat/completions?api-version=2025-01-01-preview",
    openai_api_version="2025-01-01-preview",
    deployment_name="gpt-4o-mini",
    temperature=0.3
)

qa_chain = RetrievalQA.from_chain_type(
    llm=llm,
    retriever=retriever,
    return_source_documents=False
)

st.subheader("🔎 Request a Contract Template")
contract_type = st.text_input("Enter contract type (e.g., NDA, MSA, Employment Agreement)")

if st.button("🚀 Generate Template") and contract_type:
    query = f"Create a detailed contract template for {contract_type}. Make sure it is formal, general-purpose, and does not include any party names."
    result = qa_chain.run(query)

    st.subheader("📑 Generated Contract Template")
    st.text_area("Generated Template", result, height=400)

    filename_txt = f"{contract_type.lower().replace(' ', '_')}_template.txt"
    with open(filename_txt, "w", encoding="utf-8") as f:
        f.write(result)
    with open(filename_txt, "rb") as f:
        st.download_button("📥 Download as .txt", f, file_name=filename_txt, mime="text/plain")

    filename_docx = f"{contract_type.lower().replace(' ', '_')}_template.docx"
    doc = Document()
    doc.add_heading(contract_type.upper() + " TEMPLATE", level=1)
    for para in result.split("\n"):
        doc.add_paragraph(para)
    doc.save(filename_docx)
    with open(filename_docx, "rb") as f:
        st.download_button("📄 Download as .docx", f, file_name=filename_docx, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
